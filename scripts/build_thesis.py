"""Build the StegShield bachelor thesis as a .docx following the UniBuc FMI template."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG = Path(r"D:\FACULTATE\Licenta\stegshield\outputs\figures")
OUT = Path(r"D:\FACULTATE\Licenta\StegShield_Licenta.docx")

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

def set_body_justify():
    bt = doc.styles["Body Text"]
    bt.font.name = "Times New Roman"
    bt.font.size = Pt(12)
    bt.paragraph_format.line_spacing = 1.5
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bt.paragraph_format.first_line_indent = Cm(1.0)
    bt.paragraph_format.space_after = Pt(6)
set_body_justify()

def para(text="", style=None, align=None, bold=False, size=None, italic=False, indent=True):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if style == "Body Text" and not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    return p

def body(text):
    return para(text, style="Body Text")

def h1(text):
    return doc.add_heading(text, level=1)

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def bullet(text):
    p = doc.add_paragraph(text, style="List Paragraph")
    p.style = doc.styles["List Paragraph"]
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr');
    p.paragraph_format.left_indent = Cm(1.0)
    # simple bullet via run prefix to stay portable
    p.clear()
    r = p.add_run("•  " + text)
    return p

def add_figure(filename, caption, width_cm=14.5):
    path = FIG / filename
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(10)

def add_table(headers, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(style="Normal")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(10)
    return t

# =====================================================================
# TITLE PAGE
# =====================================================================
para("UNIVERSITATEA DIN BUCUREȘTI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
para("FACULTATEA DE MATEMATICĂ ȘI INFORMATICĂ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
para("SPECIALIZAREA INFORMATICĂ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
for _ in range(3):
    para("")
para("Lucrare de licență", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17)
for _ in range(2):
    para("")
para("StegShield: Sistem hibrid pentru detecția imaginilor periculoase",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
para("prin steganaliză cu rețele neuronale convoluționale",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
para("și analiza structurală a fișierelor",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
for _ in range(4):
    para("")
para("Absolvent", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
para("Andrei-Theodor [Nume de familie]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("")
para("Coordonator științific", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
para("[Titlul și numele profesorului coordonator]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
for _ in range(4):
    para("")
para("București, iunie 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_page_break()

# =====================================================================
# REZUMAT
# =====================================================================
doc.add_heading("Rezumat", level=1)
body("Fișierele imagine sunt frecvent considerate inofensive, însă pot transporta date ascunse "
     "prin steganografie sau structuri de fișier anormale folosite în atacuri de tip stegomalware. "
     "Lucrarea de față propune StegShield, un sistem hibrid de triere a riscului care combină "
     "steganaliza vizuală bazată pe rețele neuronale convoluționale (CNN) cu analiza statică a "
     "metadatelor și a structurii fișierului, producând o etichetă finală de risc în trei trepte: "
     "sigur, suspect sau periculos.")
body("Contribuția principală nu constă într-o arhitectură CNN inedită, ci în alinierea steganalizei "
     "cu geometria reală de încorporare a atacatorului și în transformarea estimării dimensiunii "
     "payload-ului într-un semnal de severitate. Modelul propus, StegShieldCNN, extinde printr-un "
     "cap de regresie multi-task o rețea reziduală cu front-end SRM fix, estimând simultan "
     "probabilitatea de steganografie și dimensiunea payload-ului încorporat secvențial prin LSB. "
     "Pentru a evita circularitatea, capul de regresie este antrenat pe payload-uri sintetice de "
     "dimensiune cunoscută, independent de estimatorul statistic clasic.")
body("Experimentele pe setul de date Kaggle Stego Images arată o detecție binară aproape perfectă "
     "(macro-F1 = 0,9998), fără regresie la introducerea capului de regresie, o eroare medie absolută "
     "a estimării payload-ului de 62 de octeți (mediană) și un acord puternic între estimatorul CNN și "
     "cel statistic (corelație Pearson 0,977 pe scară logaritmică). Stratul de fuziune atinge un "
     "macro-F1 de 0,996 la clasificarea finală a riscului.")
para("Cuvinte cheie: ", bold=True).add_run(
    "steganaliză, rețele neuronale convoluționale, steganaliză cantitativă, stegomalware, "
    "filtre SRM, învățare multi-task, securitatea imaginilor.")
doc.add_page_break()

# =====================================================================
# ABSTRACT
# =====================================================================
doc.add_heading("Abstract", level=1)
body("Image files are commonly treated as harmless, yet they can carry data hidden through "
     "steganography or malformed file structures used in stegomalware attacks. This thesis proposes "
     "StegShield, a hybrid risk-triage system that combines CNN-based visual steganalysis with static "
     "analysis of file metadata and structure, producing a final three-level risk label: safe, "
     "suspicious or dangerous.")
body("The main contribution is not a novel CNN backbone, but the alignment of steganalysis with the "
     "attacker's real embedding geometry and the use of payload-size estimation as a severity signal. "
     "The proposed model, StegShieldCNN, augments a residual network with a fixed SRM front-end by a "
     "multi-task regression head that jointly estimates the steganography probability and the size of a "
     "sequentially LSB-embedded payload. To avoid circularity, the regression head is trained on "
     "synthetic payloads of known size, independently of the classical statistical estimator.")
body("Experiments on the Kaggle Stego Images dataset show near-perfect binary detection "
     "(macro-F1 = 0.9998) with no regression when adding the regression head, a median absolute "
     "payload estimation error of 62 bytes, and strong agreement between the CNN and statistical "
     "estimators (Pearson correlation 0.977 on a log scale). The fusion layer reaches a macro-F1 of "
     "0.996 on the final risk classification.")
para("Keywords: ", bold=True).add_run(
    "steganalysis, convolutional neural networks, quantitative steganalysis, stegomalware, "
    "SRM filters, multi-task learning, image security.")
doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (Word field; updates on open)
# =====================================================================
doc.add_heading("Cuprins", level=1)
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
t = OxmlElement('w:t'); t.text = "Faceți clic dreapta și alegeți „Update Field” pentru a genera cuprinsul."
fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar); run._r.append(instr); run._r.append(fldChar2); run._r.append(t); run._r.append(fldChar3)
doc.add_page_break()

# =====================================================================
# CH1 INTRODUCERE
# =====================================================================
h1("Introducere")

h2("Context și motivație")
body("Formatele de imagine raster precum PNG sau BMP stochează valori de pixel pe care utilizatorii "
     "le percep ca pe un simplu conținut vizual. În realitate, un fișier imagine este un container "
     "structurat care poate ascunde informație suplimentară în mai multe moduri: prin modificarea "
     "biților cel mai puțin semnificativi ai pixelilor (steganografie LSB), prin câmpuri de metadate "
     "manipulate, prin date adăugate după marcatorul de sfârșit al imaginii sau prin semnături de "
     "fișiere executabile încorporate. Aceste tehnici stau la baza atacurilor de tip stegomalware, în "
     "care un payload malițios este transportat în interiorul unui fișier aparent inofensiv, eludând "
     "filtrele care inspectează doar tipul de fișier.")
body("Detecția automată a acestor amenințări este o problemă de învățare automată dificilă, deoarece "
     "semnalul steganografic este, prin construcție, slab și greu de distins de zgomotul natural al "
     "imaginii. Steganaliza modernă în domeniul spațial folosește rețele neuronale convoluționale "
     "specializate, care amplifică reziduurile de înaltă frecvență ale imaginii înainte de clasificare. "
     "Totuși, o rețea pur vizuală nu poate detecta indicatorii de risc de la nivelul fișierului, iar o "
     "analiză statică a structurii nu poate sesiza încorporările vizuale subtile. Combinarea celor două "
     "surse de evidență reprezintă motivația centrală a acestei lucrări.")

h2("Problema abordată")
body("Lucrarea abordează detecția fișierelor imagine periculoase ca o problemă de triere a riscului în "
     "trei trepte — sigur, suspect, periculos — pe baza a două surse complementare de evidență: "
     "probabilitatea de steganografie estimată de un CNN și un set de indicatori de risc extrași prin "
     "analiză statică. Spre deosebire de literatura clasică de steganaliză, care se oprește la decizia "
     "binară „stego / curat”, sistemul propus trebuie să producă o etichetă de severitate utilă "
     "operațional și o explicație asociată.")

h2("Contribuții")
body("Contribuțiile principale ale acestei lucrări sunt următoarele:")
bullet("Un sistem hibrid de triere a riscului care integrează steganaliza CNN cu analiza statică a "
       "metadatelor și structurii fișierului într-o etichetă finală explicabilă.")
bullet("O preprocesare conștientă de geometria de încorporare: deoarece uneltele de stegomalware "
       "din familia Invoke-PSImage scriu payload-ul secvențial începând de la pixelul (0,0), imaginea "
       "este decupată din colțul stânga-sus și păstrată la valori brute 0–255, astfel încât semnalul "
       "steganografic să nu fie distrus de redimensionare sau de o decupare centrală.")
bullet("Un cap de regresie multi-task care extinde modelul StegShieldCNN pentru a estima dimensiunea "
       "payload-ului încorporat, transformând astfel steganaliza într-o problemă cantitativă.")
bullet("O metodologie anti-circularitate: capul de regresie este antrenat pe payload-uri sintetice de "
       "dimensiune cunoscută, independent de estimatorul statistic clasic, ceea ce permite o comparație "
       "validă între cele două estimatoare.")
bullet("Un strat de fuziune în care dimensiunea estimată a payload-ului — nu doar prezența "
       "steganografiei — determină severitatea, separând fișierele suspecte de cele periculoase.")

h2("Structura lucrării")
body("Capitolul 2 prezintă preliminariile teoretice privind steganografia, steganaliza spațială, "
     "modelele bogate SRM, abordările bazate pe învățare profundă și steganaliza cantitativă. "
     "Capitolul 3 descrie arhitectura sistemului propus, de la preprocesare la stratul de fuziune. "
     "Capitolul 4 detaliază metodologia experimentală, inclusiv setul de date și protocolul de "
     "antrenare. Capitolul 5 prezintă și discută rezultatele experimentale. Capitolul 6 trage "
     "concluziile și schițează direcțiile viitoare.")
doc.add_page_break()

# =====================================================================
# CH2 PRELIMINARII
# =====================================================================
h1("Preliminarii")

h2("Steganografie și steganaliză")
body("Steganografia este disciplina ascunderii informației astfel încât însăși existența mesajului să "
     "fie disimulată. În domeniul imaginilor, cea mai simplă metodă este înlocuirea biților cel mai "
     "puțin semnificativi (LSB) ai valorilor de pixel cu biții mesajului, modificare imperceptibilă "
     "vizual. Steganaliza este disciplina complementară, care urmărește să detecteze prezența unui "
     "mesaj ascuns. Steganaliza poate fi calitativă (decizia binară curat/stego) sau cantitativă, când "
     "se estimează și dimensiunea payload-ului încorporat.")

h2("Steganaliză în domeniul spațial și modele bogate (SRM)")
body("Provocarea fundamentală a steganalizei spațiale este raportul semnal-zgomot extrem de mic: "
     "perturbația introdusă de încorporare este de ordinul unui bit pe valoare de pixel, fiind ușor "
     "dominată de conținutul imaginii. Modelele bogate spațiale (Spatial Rich Models, SRM) propuse de "
     "Fridrich și Kodovský [6] rezolvă această problemă prin aplicarea unui set de filtre liniare de "
     "înaltă frecvență care suprimă conținutul imaginii și scot în evidență reziduurile de zgomot, "
     "acolo unde artefactele de încorporare sunt mai vizibile.")

h2("Steganaliză bazată pe învățare profundă")
body("Trecerea de la trăsături proiectate manual la rețele neuronale convoluționale a marcat o "
     "îmbunătățire majoră a steganalizei. Ye-Net [3] a introdus un front-end inițializat cu filtrele "
     "SRM și activarea de trunchiere (Truncated Linear Unit, TLU), care limitează amplitudinea "
     "reziduurilor. Yedroudj-Net [2] a consolidat această direcție, propunând o rețea eficientă cu un "
     "banc fix de 30 de filtre SRM, activare în valoare absolută, normalizare pe loturi și straturi de "
     "trunchiere; în această lucrare Yedroudj-Net este folosit ca referință din literatură. SRNet [4] a "
     "demonstrat eficacitatea blocurilor reziduale de tip ResNet în steganaliză, iar WISERNet [5] a "
     "tratat explicit steganaliza imaginilor color, argumentând că însumarea pe canale a convoluției "
     "obișnuite acționează ca o „coluziune liniară” ce atenuează zgomotul necorelat dintre canale — de "
     "aceea convoluția pe canale separate este preferabilă în straturile inferioare.")

h2("Steganaliză cantitativă")
body("Steganaliza cantitativă urmărește estimarea dimensiunii payload-ului, nu doar detecția "
     "prezenței acestuia. Abordările clasice se bazează pe statistici ale planului LSB; atacul "
     "histogramei și metodele structurale ale lui Westfeld și Pfitzmann [1] exploatează faptul că "
     "încorporarea secvențială lasă o porțiune inițială de tip zgomot, urmată de o tranziție abruptă "
     "către statistici specifice imaginii. Mai recent, Chen, Boroumand și Fridrich [7] au formulat "
     "estimarea payload-ului ca o problemă de regresie rezolvată cu un CNN. Lucrarea de față combină "
     "ambele perspective: un estimator statistic structural și un cap de regresie CNN, comparate între ele.")

h2("Stegomalware și analiza structurală a fișierelor")
body("Stegomalware desemnează codul malițios ascuns în fișiere media aparent inofensive [8]. Detecția "
     "sa nu se reduce la steganaliză vizuală: indicatori precum nepotrivirea dintre extensie și tipul "
     "real, datele adăugate după sfârșitul imaginii, semnăturile de executabil încorporate sau șirurile "
     "suspecte din metadate sunt esențiale. Aceste semnale de la nivelul fișierului sunt complementare "
     "evidenței vizuale și stau la baza analizorului static din sistemul propus.")
doc.add_page_break()

# =====================================================================
# CH3 ARHITECTURA
# =====================================================================
h1("Arhitectura sistemului propus")

h2("Privire de ansamblu")
body("StegShield este organizat ca o conductă de analiză cu trei componente: (1) un clasificator "
     "vizual CNN care estimează probabilitatea de steganografie și, opțional, dimensiunea payload-ului; "
     "(2) un analizor static care extrage indicatori de risc din metadate și din structura fișierului; "
     "(3) un strat de fuziune care combină cele două surse de evidență într-o etichetă finală de risc. "
     "Eticheta finală nu este atribuită de CNN singur — rețeaua estimează doar evidența vizuală de "
     "steganografie, iar decizia de severitate aparține stratului de fuziune.")

h2("Preprocesare conștientă de geometria de încorporare")
body("Steganaliza academică presupune de regulă o încorporare adaptivă, distribuită pe toată imaginea, "
     "și folosește decupări sau redimensionări arbitrare. În schimb, uneltele reale de stegomalware "
     "scriu payload-ul secvențial, în ordinea de scanare a pixelilor, începând de la pixelul (0,0). "
     "Pornind de la această observație, conducta de preprocesare a StegShield decupează colțul "
     "stânga-sus al imaginii la dimensiunea modelului, în loc să o redimensioneze, și furnizează "
     "rețelei valori de pixel brute în intervalul 0–255. Astfel, pe de o parte semnalul steganografic "
     "(concentrat în primele rânduri) nu este eliminat de o decupare centrală, iar pe de altă parte "
     "pragurile de trunchiere (T=3, T=2) își păstrează semnificația definită pe scara 0–255. O "
     "decupare centrală pe imagini de 512×512 ar elimina complet regiunea încorporată pentru "
     "payload-uri scurte.")

h2("Front-end SRM fix și trunchiere")
body("Atât modelul de referință, cât și modelul propus împart același front-end: bancul fix de 30 de "
     "filtre SRM de înaltă frecvență (nuclee 5×5 nenormalizate), aplicat pe fiecare canal de culoare "
     "separat printr-o convoluție grupată, urmat de o activare de trunchiere. Aplicarea pe canale "
     "separate, în acord cu argumentul WISERNet [5], păstrează artefactele specifice fiecărui canal în "
     "loc să le atenueze prin însumare. Ponderile front-end-ului SRM sunt înghețate, astfel încât "
     "comparația dintre arhitecturi izolează proiectarea extractorului de trăsături.")

h2("Modelul propus StegShieldCNN")
body("StegShieldCNN înlocuiește blocurile convoluționale simple din Yedroudj-Net cu patru etape de "
     "blocuri reziduale cu subeșantionare stridată (32 → 64 → 128 → 256 de canale), urmate de "
     "agregare globală prin medie și un clasificator cu regularizare prin dropout. Întrucât front-end-ul "
     "SRM este identic cu cel al modelului de referință, diferența de performanță reflectă strict "
     "efectul extractorului rezidual. Antrenarea folosește optimizatorul AdamW, programare a ratei de "
     "învățare de tip cosine annealing, pierdere cross-entropy ponderată pe clase și eșantionare "
     "echilibrată, modelul fiind selectat după macro-F1 pe setul de validare.")

h2("Capul de regresie pentru estimarea payload-ului (multi-task)")
body("Contribuția centrală la nivel de model este transformarea StegShieldCNN într-o rețea multi-task. "
     "Pe lângă capul de clasificare curat/stego, se adaugă un cap de regresie — un perceptron "
     "multistrat aplicat pe trăsăturile agregate — care estimează dimensiunea payload-ului. Ținta de "
     "regresie este log2(payload + 1), plafonată la capacitatea LSB a decupajului (24 576 de octeți "
     "pentru 256×256), deoarece rețeaua nu poate vedea payload dincolo de decupaj. Pierderea de "
     "regresie (smooth L1) se calculează doar pe eșantioanele stego cu dimensiune cunoscută, fiind "
     "mascată pentru imaginile curate; pierderea totală este suma dintre cross-entropy și termenul de "
     "regresie ponderat. Un comutator de configurare păstrează compatibilitatea cu modelele antrenate "
     "anterior, fără cap de regresie.")

h2("Generarea anti-circularitate a adevărului de bază")
body("Dacă regresorul CNN ar fi antrenat pe etichete produse de estimatorul statistic, comparația "
     "ulterioară dintre cele două ar măsura doar cât de bine imită „elevul” „profesorul”. Pentru a "
     "evita această capcană, se generează un set sintetic în care payload-uri de octeți aleatori "
     "(os.urandom, date inerte) de dimensiune cunoscută sunt încorporate secvențial începând de la "
     "pixelul (0,0), replicând schema de tip Invoke-PSImage. Dimensiunile sunt eșantionate log-uniform "
     "în intervalul [16, 24576] de octeți. Imaginile-sursă pentru antrenare/validare provin exclusiv "
     "din partiția de antrenare Kaggle, iar cele pentru test exclusiv din partiția de test, eliminând "
     "scurgerea de imagini-sursă între partiții.")

h2("Analizorul de metadate și structură")
body("Analizorul static extrage o serie de indicatori de risc: tip de fișier necunoscut, nepotrivire "
     "extensie/tip, erori de parsare, date adăugate după sfârșitul imaginii, metadate neobișnuit de "
     "mari, șiruri suspecte (de exemplu „powershell”, „base64”) și un raport anormal octeți/pixel. "
     "Detecția executabilelor încorporate validează structura antetului PE (pointerul e_lfanew către "
     "„PE\\x00\\x00”) în locul potrivirii naive a celor doi octeți „MZ”, care apărea aleator în datele "
     "comprimate ale imaginii și escalada eronat ~21% dintre fișierele curate la „periculos”. Un "
     "estimator de payload LSB secvențial, inspirat din Westfeld și Pfitzmann [1], măsoară lungimea "
     "rulării inițiale de tip zgomot din planul LSB; payload-urile la scară de script/binar (≥ 128 de "
     "octeți) au severitate ridicată.")

h2("Stratul de fuziune și etichetarea riscului")
body("Stratul de fuziune combină liniar probabilitatea de steganografie a CNN-ului și scorul de risc "
     "al metadatelor, cu reguli suplimentare de escaladare pentru evidențe de severitate ridicată "
     "(semnături de executabil încorporate, payload-uri LSB la scară de script). Scorul final este "
     "transpus în etichetele sigur / suspect / periculos. Estimarea dimensiunii payload-ului — fie "
     "statistică, fie produsă de capul de regresie CNN — alimentează indicatorii de severitate, astfel "
     "încât separarea dintre „suspect” și „periculos” să reflecte amploarea reală a payload-ului.")
doc.add_page_break()

# =====================================================================
# CH4 METODOLOGIE
# =====================================================================
h1("Metodologie experimentală")

h2("Setul de date")
body("Experimentele folosesc setul de date Kaggle Stego Images, în care payload-urile sunt încorporate "
     "prin LSB secvențial începând de la pixelul (0,0). Etichetele de risc safe/suspicious/dangerous "
     "sunt mapate la sarcina binară de steganaliză astfel: safe → clean, iar suspicious și dangerous → "
     "stego. Partiția de antrenare conține 16 000 de imagini, partiția de test standard 8 000, iar "
     "partiția adversarială 12 000 de imagini periculoase cu payload-uri codificate base64 sau "
     "arhivate ZIP. Distribuția pe clase este prezentată în Tabelul 1.")
add_table(
    ["Partiție", "safe", "suspicious", "dangerous", "Total"],
    [["train", "4 000", "4 885", "7 115", "16 000"],
     ["test_standard", "2 000", "2 481", "3 519", "8 000"],
     ["test_adversarial", "0", "0", "12 000", "12 000"]],
    caption="Tabelul 1. Distribuția pe clase a partițiilor setului de date Kaggle Stego Images."
)

h2("Protocol de antrenare")
body("Toate modelele sunt antrenate timp de 15 epoci, cu dimensiunea lotului 16, dimensiunea imaginii "
     "256×256, normalizare raw255 și decupare din colțul stânga-sus. Se folosesc optimizatorul AdamW "
     "(rată de învățare 0,001, penalizare L2 0,0001), programare cosine annealing, pierdere "
     "cross-entropy ponderată pe clase, eșantionare echilibrată și antrenare în precizie mixtă (AMP) pe "
     "GPU. Checkpoint-ul optim este selectat după macro-F1 pe validare. Pentru modelul multi-task, "
     "ponderea pierderii de regresie este 0,5.")

h2("Metrici")
body("Întrucât sarcina binară este dezechilibrată, acuratețea simplă este insuficientă: un model care "
     "prezice doar clasa majoritară poate obține o acuratețe ridicată. Raportarea pune accent pe "
     "macro-F1, acuratețe echilibrată, precizie și recall pe fiecare clasă, matricea de confuzie și "
     "numărul de fals-negative. Pentru regresia payload-ului se raportează eroarea medie absolută "
     "(MAE) și eroarea absolută mediană în octeți, precum și MAE pe scara log2. Acordul dintre "
     "estimatoare este măsurat prin corelațiile Pearson și Spearman pe scara logaritmică.")
doc.add_page_break()

# =====================================================================
# CH5 REZULTATE
# =====================================================================
h1("Rezultate experimentale și discuții")

h2("Detecția binară: StegShieldCNN vs. Yedroudj-Net")
body("Pe partiția de test standard, atât modelul de referință Yedroudj-Net, cât și modelul propus "
     "StegShieldCNN ating o detecție binară aproape perfectă (Tabelul 2). Cele două modele sunt practic "
     "la saturație pe această sarcină, ceea ce era de așteptat: încorporarea LSB la intensitate maximă "
     "pe imagini fără pierderi, cu payload pornind de la pixelul (0,0) și o decupare îndreptată exact "
     "spre acesta, constituie o problemă ușoară pentru orice rețea cu front-end SRM. Acest rezultat "
     "justifică deplasarea axei principale de evaluare către estimarea cantitativă a payload-ului, care "
     "nu poate satura precum acuratețea.")
add_table(
    ["Model", "Acuratețe", "Macro-F1", "Acuratețe echilibrată"],
    [["Yedroudj-Net (referință)", "0,99988", "0,99983", "0,99992"],
     ["StegShieldCNN (detecție)", "0,99988", "0,99983", "0,99975"],
     ["StegShieldCNN multi-task", "0,99988", "0,99983", "0,99975"]],
    caption="Tabelul 2. Detecția binară pe partiția de test standard."
)
add_figure("model_comparison.png",
           "Figura 1. Comparația acurateței, macro-F1 și acurateței echilibrate între modele.")

h2("Verificarea non-regresiei detecției în regim multi-task")
body("O întrebare esențială este dacă adăugarea capului de regresie degradează detecția. Modelul "
     "multi-task obține pe partiția de test standard exact aceleași valori ca modelul de detecție pură "
     "(macro-F1 = 0,99983), deci capul de regresie nu produce nicio regresie a detecției. Curbele de "
     "antrenare din Figura 2 confirmă o convergență stabilă a ambelor sarcini, eroarea medie absolută a "
     "payload-ului pe validare coborând la aproximativ 537 de octeți.")
add_figure("steganalysis_multitask_training_training_curves.png",
           "Figura 2. Curbele de antrenare ale modelului multi-task (pierdere și macro-F1 pe validare).")

h2("Calitatea regresiei payload-ului")
body("Pe setul de test sintetic (1 601 de eșantioane supervizate, capacitate 24 576 de octeți), "
     "regresorul CNN obține o eroare absolută mediană de doar 62 de octeți și o eroare medie pe scara "
     "log2 de 0,20. Eroarea medie absolută de 533 de octeți este dominată de un număr mic de "
     "payload-uri mari, ceea ce confirmă alegerea unei ținte logaritmice. Figura 3 prezintă dispersia "
     "valorilor estimate față de cele reale, aliniate strâns pe diagonală pe trei ordine de mărime.")
add_table(
    ["Metrică", "Valoare"],
    [["Eroare absolută mediană", "62 octeți"],
     ["Eroare medie absolută (MAE)", "533,8 octeți"],
     ["MAE pe scara log2", "0,20"],
     ["Eșantioane supervizate", "1 601"],
     ["Capacitate decupaj", "24 576 octeți"]],
    caption="Tabelul 3. Calitatea regresiei payload-ului pe setul de test sintetic."
)
add_figure("payload_regression_test_scatter.png",
           "Figura 3. Dimensiunea payload-ului estimată de CNN față de cea reală (scară log-log).")

h2("Acordul dintre estimatorul CNN și cel statistic")
body("Pe imaginile stego reale din partiția de test standard (5 996 de eșantioane comparabile), "
     "estimatorul CNN și estimatorul statistic Westfeld–Pfitzmann prezintă un acord puternic: "
     "corelație Pearson de 0,977 și corelație Spearman de 0,940 pe scara logaritmică. Întrucât cele "
     "două estimatoare au fost dezvoltate independent — CNN-ul fiind antrenat pe adevăr de bază "
     "sintetic, nu pe ieșirile estimatorului statistic — acest acord constituie o validare reciprocă "
     "veritabilă, nu o consecință a circularității. Diferența absolută mediană de 891 de octeți "
     "reflectă rezoluția în blocuri a estimatorului statistic.")
add_figure("payload_agreement_test_standard_scatter.png",
           "Figura 4. Acordul dintre estimarea CNN și cea statistică a payload-ului (scară log-log).")

h2("Evaluarea stratului de fuziune")
body("Tabelul 4 compară clasificarea finală a riscului pentru cele trei surse de severitate a "
     "payload-ului. Sursa statistică oferă cel mai bun compromis (macro-F1 = 0,996, recall pentru clasa "
     "„periculos” = 0,992). Sursa CNN obține rezultate apropiate (macro-F1 = 0,982). Combinarea ambelor "
     "surse („both”) duce însă la supra-escaladare: ambele indicatoare de payload se declanșează "
     "simultan, multe eșantioane „suspect” fiind clasificate drept „periculos”, ceea ce prăbușește "
     "macro-F1 la 0,580, deși recall-ul pentru clasa „periculos” atinge 1,0. Concluzia practică este că "
     "o singură sursă de severitate, calibrată corespunzător, este preferabilă combinării naive a "
     "ambelor. Atât baza metadata-only, cât și baza cnn-only nu pot produce singure eticheta "
     "„periculos”, ceea ce justifică necesitatea fuziunii.")
add_table(
    ["Sursă payload", "Acuratețe", "Macro-F1", "Recall „periculos”"],
    [["statistical", "0,99575", "0,99604", "0,992"],
     ["cnn", "0,98050", "0,98225", "0,956"],
     ["both", "0,68963", "0,57962", "1,000"]],
    caption="Tabelul 4. Clasificarea finală a riscului prin fuziune, pe cele trei surse de severitate."
)
add_figure("fusion_test_standard_metrics_fused_confusion_matrix.png",
           "Figura 5. Matricea de confuzie a clasificatorului fuzionat (sursă statistică).")

h2("Limitări")
body("Rezultatele trebuie interpretate cu prudență. Evaluarea folosește un singur set de date și o "
     "singură familie de încorporare (LSB secvențial), iar partiția de test standard este saturată "
     "pentru detecția binară. Sistemul nu acoperă steganaliza în domeniul JPEG și nici încorporările "
     "adaptive la conținut (S-UNIWARD, WOW). Estimatorul statistic presupune un flux de biți "
     "Bernoulli(0,5), ipoteză care nu se aplică payload-urilor codificate base64 din partiția "
     "adversarială. În fine, ponderile și pragurile stratului de fuziune sunt calibrate manual; o "
     "fuziune învățată ar constitui o extensie firească.")
doc.add_page_break()

# =====================================================================
# CH6 CONCLUZII
# =====================================================================
h1("Concluzii și direcții viitoare")
body("Lucrarea a propus StegShield, un sistem hibrid de triere a riscului fișierelor imagine care "
     "combină steganaliza CNN cu analiza statică a structurii fișierului. Contribuția distinctivă nu "
     "este o nouă arhitectură de rețea, ci alinierea steganalizei cu geometria reală de încorporare a "
     "atacatorului și utilizarea dimensiunii payload-ului ca semnal de severitate, validată printr-o "
     "metodologie experimentală riguroasă și non-circulară.")
body("Rezultatele arată o detecție binară aproape perfectă, fără regresie la introducerea capului de "
     "regresie multi-task, o estimare precisă a dimensiunii payload-ului (eroare mediană de 62 de "
     "octeți) și un acord puternic între estimatorul CNN și cel statistic (Pearson 0,977). Stratul de "
     "fuziune atinge un macro-F1 de 0,996 la clasificarea finală a riscului, sursa statistică de "
     "severitate dovedindu-se cea mai robustă.")
body("Direcțiile viitoare includ: evaluarea pe seturi de date suplimentare și pe familii de "
     "încorporare adaptive (S-UNIWARD, WOW); extinderea către steganaliza în domeniul JPEG; "
     "antrenarea cu mai multe semințe aleatoare și raportarea intervalelor de încredere; un studiu de "
     "ablație pentru decuparea stânga-sus față de cea centrală, care ar evidenția direct importanța "
     "geometriei de încorporare; și înlocuirea regulilor manuale de fuziune cu un clasificator învățat "
     "(de exemplu regresie logistică) peste trăsăturile combinate.")
doc.add_page_break()

# =====================================================================
# BIBLIOGRAFIE
# =====================================================================
h1("Bibliografie")
refs = [
 '[1] A. Westfeld, A. Pfitzmann, „Attacks on Steganographic Systems”, în Information Hiding 1999, '
 'LNCS 1768, pp. 61–76.',
 '[2] M. Yedroudj, F. Comby, M. Chaumont, „Yedroudj-Net: An Efficient CNN for Spatial Steganalysis”, '
 'IEEE ICASSP 2018, pp. 2092–2096, doi:10.1109/ICASSP.2018.8461438.',
 '[3] J. Ye, J. Ni, Y. Yi, „Deep Learning Hierarchical Representations for Image Steganalysis”, '
 'IEEE Transactions on Information Forensics and Security, vol. 12, nr. 11, 2017, pp. 2545–2557.',
 '[4] M. Boroumand, M. Chaumont, J. Fridrich, „Deep Residual Network for Steganalysis of Digital '
 'Images”, IEEE Transactions on Information Forensics and Security, vol. 14, nr. 5, 2019, pp. 1181–1193.',
 '[5] J. Zeng, S. Tan, G. Liu, B. Li, J. Huang, „WISERNet: Wider Separate-then-reunion Network for '
 'Steganalysis of Color Images”, IEEE Transactions on Information Forensics and Security, vol. 14, '
 'nr. 10, 2019, pp. 2735–2748.',
 '[6] J. Fridrich, J. Kodovský, „Rich Models for Steganalysis of Digital Images”, IEEE Transactions '
 'on Information Forensics and Security, vol. 7, nr. 3, 2012, pp. 868–882.',
 '[7] M. Chen, M. Boroumand, J. Fridrich, „Deep Learning Regressors for Quantitative Steganalysis”, '
 'în Media Watermarking, Security, and Forensics, IS&T Electronic Imaging, 2018.',
 '[8] D. Puchalski et al., „Stegomalware Detection through Structural Analysis of Media Files”, în '
 'Proceedings of the 15th International Conference on Availability, Reliability and Security (ARES), '
 '2020, doi:10.1145/3407023.3409187.',
 '[9] A. Paszke et al., „PyTorch: An Imperative Style, High-Performance Deep Learning Library”, '
 'Advances in Neural Information Processing Systems 32 (NeurIPS), 2019.',
 '[10] Kaggle Stego Images Dataset, https://www.kaggle.com/datasets/marcozuppelli/stegoimagesdataset.',
]
for r in refs:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(r)
    run.font.size = Pt(11)

doc.save(str(OUT))
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs))
